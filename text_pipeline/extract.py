"""
Text Pipeline — 文本提取 & 编码检测

extract_text / detect_encoding / detect_language。
"""

import os
import re
import logging

from docx import Document
from bs4 import BeautifulSoup

from .ocr import ocr_image

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════
# 编码检测
# ═══════════════════════════════════════════

def detect_encoding(file_path: str, sample_size: int = 10000) -> str:
    """
    检测文件编码。优先 chardet，失败后用 UTF-8 → GBK → latin-1 兜底链。
    sample_size: 用于检测的字节数（默认 10000，约 10KB）
    """
    with open(file_path, "rb") as f:
        raw = f.read(sample_size)
    if not raw:
        return "utf-8"  # 空文件，默认 UTF-8

    # 先试 chardet（如果已安装）
    try:
        import chardet
        result = chardet.detect(raw)
        enc = result.get("encoding", "").strip().lower()
        conf = result.get("confidence", 0)
        if enc and conf >= 0.6:
            enc_map = {
                "utf-8": "utf-8",
                "ascii": "utf-8",
                "gb2312": "gbk",
                "gbk": "gbk",
                "gb18030": "gb18030",
                "big5": "big5",
                "iso-8859-1": "latin-1",
                "windows-1252": "cp1252",
            }
            return enc_map.get(enc, enc)
    except ImportError:
        pass  # chardet 未安装，走兜底链

    # 兜底链：UTF-8 → GBK → latin-1
    try:
        raw.decode("utf-8")
        return "utf-8"
    except UnicodeDecodeError:
        pass
    try:
        raw.decode("gbk")
        return "gbk"
    except UnicodeDecodeError:
        pass
    # 最后兜底：latin-1 永不失败（但可能乱码）
    return "latin-1"


# ═══════════════════════════════════════════
# 语言检测
# ═══════════════════════════════════════════

def _detect_language(text: str) -> str:
    """通过 Unicode 区块统计检测语言（前 2000 字）。"""
    sample = text[:2000]
    if not sample.strip():
        return "zh"
    total = len(sample)
    cjk = sum(1 for c in sample if '\u4e00' <= c <= '\u9fff' or '\u3400' <= c <= '\u4dbf')
    hiragana = sum(1 for c in sample if '\u3040' <= c <= '\u309f')
    katakana = sum(1 for c in sample if '\u30a0' <= c <= '\u30ff')
    hangul = sum(1 for c in sample if '\uac00' <= c <= '\ud7af')
    latin = sum(1 for c in sample if c.isascii() and c.isalpha())

    cjk_ratio = cjk / total
    ja_ratio = (hiragana + katakana) / total
    ko_ratio = hangul / total
    en_ratio = latin / total

    if cjk_ratio >= 0.30:
        return "zh"
    if ja_ratio >= 0.10:
        return "ja"
    if ko_ratio >= 0.10:
        return "ko"
    if en_ratio >= 0.60:
        return "en"
    return "zh"  # 兜底


def detect_language(text: str) -> str:
    """
    程序检测文本语言（中/英），不调用 LLM，确定性输出。

    逻辑：
        - 统计 CJK 统一汉字范围（\u4e00-\u9fff）字符占比
        - 占比 > 30% → "zh"
        - 否则 → "en"（默认英文）
        - 空文本 → "en"

    返回:
        "zh" | "en" | "ja" | "ko"  (远期可扩展)
    """
    if not text:
        return "en"
    total = len(text)
    if total == 0:
        return "en"
    cjk_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    if cjk_count / total > 0.3:
        return "zh"
    return "en"


# ═══════════════════════════════════════════
# 文本提取
# ═══════════════════════════════════════════

def extract_text(file_path: str) -> dict:
    """
    统一文本提取入口：根据文件扩展名调用对应解析器。

    支持格式：
      - .txt / .md / .json / .csv → 直接读取（自动检测编码）
      - .docx → python-docx 提取文本
      - .html / .htm → BeautifulSoup 提取文本（去除标签）
      - .srt → 解析 SRT 字幕格式（去除时间戳）
      - .pdf → 尝试 pdfplumber 提取文本

    返回:
        {"ok": True, "text": "...", "chars": N, "meta": {}}
        {"ok": False, "error": "..."}
    """
    if not os.path.exists(file_path):
        return {"ok": False, "error": f"文件不存在: {file_path}"}

    ext = os.path.splitext(file_path)[1].lower()

    # ── 纯文本格式：直接读取 ──
    if ext in (".txt", ".md", ".json", ".csv", ".log"):
        encoding = detect_encoding(file_path)
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read()
            return {"ok": True, "text": text, "chars": len(text), "meta": {"encoding": encoding}}
        except Exception as e:
            return {"ok": False, "error": f"读取文件失败: {e}"}

    # ── DOCX 格式 ──
    if ext == ".docx":
        try:
            doc = Document(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            text = "\n\n".join(parts)
            return {"ok": True, "text": text, "chars": len(text), "meta": {"format": "docx"}}
        except Exception as e:
            return {"ok": False, "error": f"解析 DOCX 失败: {e}"}

    # ── HTML 格式 ──
    if ext in (".html", ".htm"):
        try:
            with open(file_path, "r", encoding=detect_encoding(file_path)) as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            for tag in soup(["script", "style"]):
                tag.decompose()
            text = soup.get_text(separator="\n")
            text = re.sub(r"\n{3,}", "\n\n", text).strip()
            return {"ok": True, "text": text, "chars": len(text), "meta": {"format": "html"}}
        except Exception as e:
            return {"ok": False, "error": f"解析 HTML 失败: {e}"}

    # ── SRT 字幕格式 ──
    if ext == ".srt":
        try:
            encoding = detect_encoding(file_path)
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            lines = []
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.isdigit():
                    continue
                if "-->" in line:
                    continue
                lines.append(line)
            text = "\n".join(lines)
            return {"ok": True, "text": text, "chars": len(text), "meta": {"format": "srt", "encoding": encoding}}
        except Exception as e:
            return {"ok": False, "error": f"解析 SRT 失败: {e}"}

    # ── PDF 格式（混合模式：先提取文本，失败则用 OCR）──
    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        parts.append(page_text)
                text = "\n\n".join(parts)

            # 如果提取到足够文本，直接返回
            if text.strip() and len(text.strip()) > 100:
                return {"ok": True, "text": text, "chars": len(text), "meta": {"format": "pdf", "pages": len(pdf.pages), "ocr": False}}

            # 文本太少，尝试 OCR（逐页转图片后识别）
            logger.info(f"[PDF] 文本提取不足，启用 OCR: {file_path}")
            ocr_parts = []
            page_count = 0
            try:
                from pdf2image import convert_from_path
                images = convert_from_path(file_path, dpi=200)
                page_count = len(images)
                for i, img in enumerate(images):
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                        img.save(tmp.name, "PNG")
                        result = ocr_image(tmp.name)
                        if result.get("ok") and result.get("text"):
                            ocr_parts.append(f"--- 第 {i+1} 页 ---\n{result['text']}")
                        os.unlink(tmp.name)
            except ImportError:
                # pdf2image 未安装，尝试用 pymupdf
                try:
                    import fitz  # pymupdf
                    doc = fitz.open(file_path)
                    page_count = len(doc)
                    for i, page in enumerate(doc):
                        pix = page.get_pixmap(dpi=200)
                        import tempfile
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                            pix.save(tmp.name)
                            result = ocr_image(tmp.name)
                            if result.get("ok") and result.get("text"):
                                ocr_parts.append(f"--- 第 {i+1} 页 ---\n{result['text']}")
                            os.unlink(tmp.name)
                except ImportError:
                    return {"ok": False, "error": "PDF 无文本内容，且未安装 OCR 依赖（pdf2image 或 pymupdf）"}

            if ocr_parts:
                text = "\n\n".join(ocr_parts)
                return {"ok": True, "text": text, "chars": len(text), "meta": {"format": "pdf", "pages": page_count, "ocr": True}}
            return {"ok": False, "error": "PDF 无文本内容，OCR 也失败"}

        except ImportError:
            return {"ok": False, "error": "需要安装 pdfplumber：pip install pdfplumber"}
        except Exception as e:
            return {"ok": False, "error": f"解析 PDF 失败: {e}"}

    # ── 图片格式：调用 OCR ──
    if ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"):
        result = ocr_image(file_path)
        if not result.get("ok"):
            return {"ok": False, "error": result.get("error", "OCR 识别失败")}
        text = result.get("text", "")
        if not text.strip():
            return {"ok": False, "error": "OCR 识别结果为空"}
        return {
            "ok": True,
            "text": text,
            "chars": len(text),
            "meta": {
                "format": "image",
                "ocr_model": result.get("model", "unknown"),
                "ocr_conf": result.get("conf", 0.0),
            }
        }

    # ── 不支持的格式 ──
    return {"ok": False, "error": f"不支持的文件格式: {ext}"}
